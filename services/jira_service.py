# services/jira_service.py
import requests
import json
import openai
import logging
from typing import Dict, List, Optional, Any
from config.settings import Config
from models.jira_models import JiraQuery, QueryType, JiraIssue
import re
import uuid
import tempfile
from datetime import datetime
from docx import Document
from docx.shared import Inches
from io import BytesIO

logger = logging.getLogger(__name__)

class JiraService:
    """Service for Jira operations"""
    
    def __init__(self):
        self.jira_url = Config.JIRA_URL.rstrip('/')
        self.auth = (Config.JIRA_USERNAME, Config.JIRA_TOKEN)
        self.headers = {
            'Accept': 'application/json',
            'Content-Type': 'application/json'
        }
        openai.api_key = Config.OPENAI_API_KEY
    
    def parse_natural_language_query(self, user_query: str) -> JiraQuery:
        """Use AI to parse natural language queries into structured JiraQuery objects"""
        prompt = f"""
        Parse the following user query about Jira tickets and extract relevant information:
        
        User Query: "{user_query}"
        
        Extract and return a JSON object with the following fields (set to null if not mentioned):
        - query_type: one of ["story_search", "epic_search", "status_filter", "assignee_filter", "date_range", "project_filter"]
        - project_key: project abbreviation (e.g., "PROJ", "DEV")
        - assignee: person's name or username
        - status: ticket status (e.g., "To Do", "In Progress", "Done")
        - epic_key: specific epic identifier
        - story_key: specific story identifier
        - date_from: start date in YYYY-MM-DD format
        - date_to: end date in YYYY-MM-DD format
        - keywords: array of relevant search terms
        
        Examples:
        - "Show me all stories assigned to John" -> {{"query_type": "assignee_filter", "assignee": "John"}}
        - "Find epic PROJ-123" -> {{"query_type": "epic_search", "epic_key": "PROJ-123"}}
        - "Stories in progress for project DEV" -> {{"query_type": "status_filter", "project_key": "DEV", "status": "In Progress"}}
        
        Return only the JSON object:
        """
        
        try:
            response = openai.ChatCompletion.create(
                model=Config.OPENAI_MODEL,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1
            )
            
            parsed_data = json.loads(response.choices[0].message.content)
            
            return JiraQuery(
                query_type=QueryType(parsed_data.get('query_type', 'story_search')),
                project_key=parsed_data.get('project_key'),
                assignee=parsed_data.get('assignee'),
                status=parsed_data.get('status'),
                epic_key=parsed_data.get('epic_key'),
                story_key=parsed_data.get('story_key'),
                date_from=parsed_data.get('date_from'),
                date_to=parsed_data.get('date_to'),
                keywords=parsed_data.get('keywords', [])
            )
        except Exception as e:
            logger.error(f"Error parsing query: {e}")
            return JiraQuery(
                query_type=QueryType.STORY_SEARCH,
                keywords=user_query.split()
            )
    
    def build_jql_query(self, parsed_query: JiraQuery) -> str:
        """Convert parsed query into JQL (Jira Query Language)"""
        jql_parts = []
        
        if parsed_query.project_key:
            jql_parts.append(f'project = "{parsed_query.project_key}"')
        
        if parsed_query.query_type == QueryType.EPIC_SEARCH:
            jql_parts.append('issuetype = Epic')
        elif parsed_query.query_type == QueryType.STORY_SEARCH:
            jql_parts.append('issuetype = Story')
        
        if parsed_query.epic_key:
            jql_parts.append(f'key = "{parsed_query.epic_key}"')
        if parsed_query.story_key:
            jql_parts.append(f'key = "{parsed_query.story_key}"')
        
        if parsed_query.assignee:
            jql_parts.append(f'assignee ~ "{parsed_query.assignee}"')
        
        if parsed_query.status:
            jql_parts.append(f'status = "{parsed_query.status}"')
        
        if parsed_query.date_from:
            jql_parts.append(f'created >= "{parsed_query.date_from}"')
        if parsed_query.date_to:
            jql_parts.append(f'created <= "{parsed_query.date_to}"')
        
        if parsed_query.keywords:
            keyword_search = ' OR '.join([f'summary ~ "{kw}" OR description ~ "{kw}"' for kw in parsed_query.keywords])
            jql_parts.append(f'({keyword_search})')
        
        return ' AND '.join(jql_parts) if jql_parts else 'project is not EMPTY'
    
    def search_jira_issues(self, jql_query: str, max_results: int = 50) -> Dict[str, Any]:
        """Execute JQL query against Jira API"""
        search_url = f"{self.jira_url}/rest/api/3/search"
        
        payload = {
            'jql': jql_query,
            'maxResults': max_results,
            'fields': [
                'summary', 'description', 'status', 'assignee', 'reporter',
                'created', 'updated', 'priority', 'issuetype', 'project',
                'parent', 'subtasks', 'labels', 'components'
            ]
        }
        
        try:
            response = requests.post(
                search_url,
                headers=self.headers,
                auth=self.auth,
                data=json.dumps(payload),
                timeout=30
            )
            response.raise_for_status()
            return response.json()
        except requests.exceptions.RequestException as e:
            logger.error(f"Error querying Jira: {e}")
            raise Exception(f"Failed to query Jira: {str(e)}")
    
    def process_user_query(self, user_query: str) -> Dict[str, Any]:
        """Main method to process user queries end-to-end"""
        logger.info(f"Processing query: '{user_query}'")
        
        try:
            # Parse natural language query
            parsed_query = self.parse_natural_language_query(user_query)
            logger.info(f"Parsed query: {parsed_query}")
            
            # Build JQL query
            jql_query = self.build_jql_query(parsed_query)
            logger.info(f"JQL Query: {jql_query}")
            
            # Execute search
            issues_data = self.search_jira_issues(jql_query)
            
            return {
                'success': True,
                'data': issues_data,
                'jql_query': jql_query,
                'parsed_query': parsed_query.to_dict()
            }
        except Exception as e:
            logger.error(f"Error processing query: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def get_issue_details(self, issue_key: str) -> Dict[str, Any]:
        """Get detailed information about a specific Jira issue"""
        issue_url = f"{self.jira_url}/rest/api/3/issue/{issue_key}"
        
        params = {
            'fields': 'summary,description,status,assignee,priority,issuetype,labels,components,reporter,created,updated'
        }
        
        try:
            response = requests.get(
                issue_url,
                headers=self.headers,
                auth=self.auth,
                params=params,
                timeout=30
            )
            response.raise_for_status()
            return response.json()
        except requests.exceptions.RequestException as e:
            logger.error(f"Error fetching Jira issue {issue_key}: {e}")
            raise Exception(f"Failed to fetch issue {issue_key}: {str(e)}")
    
    def test_connection(self) -> Dict[str, Any]:
        """Test Jira connection"""
        try:
            test_jql = "project is not EMPTY ORDER BY created DESC"
            search_url = f"{self.jira_url}/rest/api/3/search"
            
            payload = {
                'jql': test_jql,
                'maxResults': 1,
                'fields': ['summary', 'status', 'issuetype', 'project']
            }
            
            response = requests.post(
                search_url,
                headers=self.headers,
                auth=self.auth,
                data=json.dumps(payload),
                timeout=10
            )
            response.raise_for_status()
            result = response.json()
            
            return {
                'success': True,
                'message': 'Jira connection successful',
                'total_issues': result.get('total', 0)
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Jira connection failed: {str(e)}'
            }

    def parse_security_command(self, user_input: str) -> Dict[str, Any]:
        """Parse natural language security analysis commands"""
        command_patterns = {
            'fraud_security_analysis': r'(?:do|perform|run)\s+fraud\s*[&+]\s*security\s+analysis\s+for\s+(?:this\s+)?ticket\s+([A-Z]+-\d+)',
            'security_analysis': r'(?:do|perform|run)\s+security\s+analysis\s+(?:for\s+)?(?:ticket\s+)?([A-Z]+-\d+)',
            'fraud_analysis': r'(?:do|perform|run)\s+fraud\s+analysis\s+(?:for\s+)?(?:ticket\s+)?([A-Z]+-\d+)',
            'analyze_ticket': r'analyze\s+(?:ticket\s+)?([A-Z]+-\d+)\s+(?:for\s+)?(?:fraud|security)'
        }
        
        user_input = user_input.strip()
        
        for intent, pattern in command_patterns.items():
            match = re.search(pattern, user_input, re.IGNORECASE)
            if match:
                ticket_key = match.group(1).upper()
                return {
                    'intent': intent,
                    'ticket_key': ticket_key,
                    'original_command': user_input,
                    'success': True
                }
        
        return {
            'intent': 'unknown',
            'error': 'Command not recognized. Try: "Do fraud & security analysis for ticket TICKET-123"',
            'success': False
        }
    
    def get_comprehensive_ticket_data(self, issue_key: str) -> Dict[str, Any]:
        """Get comprehensive ticket data for security analysis"""
        issue_url = f"{self.jira_url}/rest/api/3/issue/{issue_key}"
        
        # Expand to get more detailed information
        params = {
            'expand': 'changelog,renderedFields,names,schema,operations,editmeta,changelog,transitions',
            'fields': '*all'
        }
        
        try:
            response = requests.get(
                issue_url,
                headers=self.headers,
                auth=self.auth,
                params=params,
                timeout=30
            )
            response.raise_for_status()
            issue_data = response.json()
            
            # Get comments separately for better formatting
            comments_url = f"{self.jira_url}/rest/api/3/issue/{issue_key}/comment"
            comments_response = requests.get(
                comments_url,
                headers=self.headers,
                auth=self.auth,
                timeout=30
            )
            comments_data = comments_response.json() if comments_response.status_code == 200 else {'comments': []}
            
            # Structure the comprehensive data
            fields = issue_data.get('fields', {})
            
            comprehensive_data = {
                'key': issue_data.get('key'),
                'summary': fields.get('summary', ''),
                'description': fields.get('description', ''),
                'issue_type': fields.get('issuetype', {}).get('name', '') if fields.get('issuetype') else '',
                'priority': fields.get('priority', {}).get('name', '') if fields.get('priority') else 'None',
                'status': fields.get('status', {}).get('name', '') if fields.get('status') else '',
                'assignee': fields.get('assignee', {}).get('displayName', 'Unassigned') if fields.get('assignee') else 'Unassigned',
                'reporter': fields.get('reporter', {}).get('displayName', '') if fields.get('reporter') else '',
                'created': fields.get('created', ''),
                'updated': fields.get('updated', ''),
                'resolved': fields.get('resolutiondate', ''),
                'labels': fields.get('labels', []),
                'components': [comp.get('name', '') for comp in fields.get('components', [])],
                'fix_versions': [v.get('name', '') for v in fields.get('fixVersions', [])],
                'affects_versions': [v.get('name', '') for v in fields.get('versions', [])],
                'project': fields.get('project', {}).get('name', '') if fields.get('project') else '',
                'environment': fields.get('environment', ''),
                'security_level': fields.get('security', {}).get('name', 'None') if fields.get('security') else 'None',
                
                # Comments with detailed info
                'comments': [{
                    'author': comment.get('author', {}).get('displayName', ''),
                    'body': comment.get('body', ''),
                    'created': comment.get('created', ''),
                    'updated': comment.get('updated', '')
                } for comment in comments_data.get('comments', [])],
                
                # Attachments
                'attachments': [{
                    'filename': att.get('filename', ''),
                    'size': att.get('size', 0),
                    'created': att.get('created', ''),
                    'author': att.get('author', {}).get('displayName', '')
                } for att in fields.get('attachment', [])],
                
                # Change history for pattern analysis
                'changelog_entries': len(issue_data.get('changelog', {}).get('histories', [])),
                'total_time_spent': fields.get('timespent', 0),
                'original_estimate': fields.get('timeoriginalestimate', 0),
                
                # Custom fields (if any)
                'custom_fields': {k: v for k, v in fields.items() if k.startswith('customfield_')},
                
                # Metadata
                'raw_data_available': True,
                'analysis_timestamp': datetime.now().isoformat()
            }
            
            return comprehensive_data
            
        except requests.exceptions.RequestException as e:
            logger.error(f"Error fetching comprehensive ticket data for {issue_key}: {e}")
            return None
    
    def perform_ai_security_analysis(self, ticket_data: Dict[str, Any]) -> Dict[str, Any]:
        """Perform AI-powered fraud and security analysis"""
        
        # Prepare context for AI analysis
        analysis_context = self._prepare_analysis_context(ticket_data)
        
        analysis_prompt = f"""
        You are an expert cybersecurity and fraud analyst. Perform a comprehensive security analysis of this Jira ticket.
        
        TICKET INFORMATION:
        {analysis_context}
        
        Please provide a detailed analysis covering:
        
        1. FRAUD RISK ASSESSMENT (Score 1-10, where 10 is highest risk):
           - Analyze description and comments for fraud indicators
           - Check for suspicious patterns, unusual requests, or social engineering attempts
           - Evaluate user behavior patterns from change history
        
        2. SECURITY VULNERABILITY ANALYSIS:
           - Identify potential security vulnerabilities mentioned
           - Assess data exposure risks
           - Check for compliance issues (PCI, GDPR, SOX, etc.)
        
        3. IMPACT ASSESSMENT:
           - Business impact level (Low/Medium/High/Critical)
           - Affected systems and data types
           - Potential financial or reputational damage
        
        4. RISK INDICATORS:
           - Red flags identified in content
           - Unusual patterns or behaviors
           - Timeline anomalies
        
        5. RECOMMENDATIONS:
           - Immediate actions required
           - Long-term security improvements
           - Monitoring recommendations
        
        6. COMPLIANCE CONCERNS:
           - Regulatory requirements affected
           - Data privacy implications
           - Audit trail requirements
        
        Format your response as structured JSON with the following keys:
        - fraud_risk_score (1-10)
        - security_risk_level (Low/Medium/High/Critical)
        - business_impact (Low/Medium/High/Critical)
        - key_findings (array of strings)
        - red_flags (array of strings)
        - recommendations (array of strings)
        - compliance_issues (array of strings)
        - executive_summary (string)
        - detailed_analysis (string)
        """
        
        try:
            response = openai.ChatCompletion.create(
                model=Config.OPENAI_MODEL or "gpt-4",
                messages=[
                    {"role": "system", "content": "You are an expert cybersecurity and fraud analyst with extensive experience in risk assessment and compliance."},
                    {"role": "user", "content": analysis_prompt}
                ],
                temperature=0.2,
                max_tokens=2500
            )
            
            analysis_text = response.choices[0].message.content
            
            # Try to parse as JSON, fallback to structured text
            try:
                analysis_result = json.loads(analysis_text)
            except json.JSONDecodeError:
                # If not valid JSON, create structured response from text
                analysis_result = self._parse_analysis_text(analysis_text)
            
            # Add metadata
            analysis_result.update({
                'ticket_key': ticket_data['key'],
                'analysis_timestamp': datetime.now().isoformat(),
                'analyst': 'AI Security Analyst',
                'analysis_version': '1.0'
            })
            
            return analysis_result
            
        except Exception as e:
            logger.error(f"Error in AI security analysis: {e}")
            return {
                'error': str(e),
                'ticket_key': ticket_data['key'],
                'fraud_risk_score': 'Unknown',
                'security_risk_level': 'Unknown',
                'executive_summary': f'Analysis failed due to error: {str(e)}'
            }
    
    def _prepare_analysis_context(self, ticket_data: Dict[str, Any]) -> str:
        """Prepare structured context for AI analysis"""
        context_parts = [
            f"Ticket Key: {ticket_data['key']}",
            f"Summary: {ticket_data['summary']}",
            f"Description: {ticket_data['description'][:1000]}..." if len(ticket_data['description']) > 1000 else f"Description: {ticket_data['description']}",
            f"Type: {ticket_data['issue_type']}",
            f"Priority: {ticket_data['priority']}",
            f"Status: {ticket_data['status']}",
            f"Assignee: {ticket_data['assignee']}",
            f"Reporter: {ticket_data['reporter']}",
            f"Created: {ticket_data['created']}",
            f"Last Updated: {ticket_data['updated']}",
            f"Security Level: {ticket_data['security_level']}",
            f"Project: {ticket_data['project']}",
            f"Components: {', '.join(ticket_data['components'])}",
            f"Labels: {', '.join(ticket_data['labels'])}",
        ]
        
        if ticket_data['comments']:
            context_parts.append(f"\nComments ({len(ticket_data['comments'])}):")
            for i, comment in enumerate(ticket_data['comments'][:5]):  # Limit to first 5 comments
                context_parts.append(f"  {i+1}. {comment['author']} ({comment['created']}): {comment['body'][:200]}...")
        
        if ticket_data['attachments']:
            context_parts.append(f"\nAttachments ({len(ticket_data['attachments'])}):")
            for att in ticket_data['attachments']:
                context_parts.append(f"  - {att['filename']} ({att['size']} bytes)")
        
        return '\n'.join(context_parts)
    
    def _parse_analysis_text(self, analysis_text: str) -> Dict[str, Any]:
        """Parse unstructured analysis text into structured format"""
        # Simple parsing fallback if JSON parsing fails
        lines = analysis_text.split('\n')
        
        result = {
            'fraud_risk_score': 'Unknown',
            'security_risk_level': 'Medium',
            'business_impact': 'Medium',
            'key_findings': [],
            'red_flags': [],
            'recommendations': [],
            'compliance_issues': [],
            'executive_summary': 'Analysis completed',
            'detailed_analysis': analysis_text
        }
        
        current_section = None
        for line in lines:
            line = line.strip()
            if 'fraud risk' in line.lower() and any(char.isdigit() for char in line):
                # Extract risk score
                import re
                numbers = re.findall(r'\d+', line)
                if numbers:
                    result['fraud_risk_score'] = numbers[0]
            elif 'security risk' in line.lower():
                if 'high' in line.lower():
                    result['security_risk_level'] = 'High'
                elif 'critical' in line.lower():
                    result['security_risk_level'] = 'Critical'
                elif 'low' in line.lower():
                    result['security_risk_level'] = 'Low'
            elif line.startswith('- ') or line.startswith('• '):
                # Add to current section
                item = line[2:].strip()
                if 'recommendation' in (current_section or '').lower():
                    result['recommendations'].append(item)
                elif 'finding' in (current_section or '').lower():
                    result['key_findings'].append(item)
                elif 'flag' in (current_section or '').lower():
                    result['red_flags'].append(item)
            
            # Track current section
            if any(keyword in line.lower() for keyword in ['recommendation', 'finding', 'flag', 'compliance']):
                current_section = line
        
        return result
    
    def generate_security_report_document(self, ticket_data: Dict[str, Any], analysis_results: Dict[str, Any]) -> bytes:
        """Generate Word document for security analysis report"""
        
        doc = Document()
        
        # Title page
        title = doc.add_heading('Fraud & Security Analysis Report', 0)
        title.alignment = 1  # Center alignment
        
        doc.add_paragraph(f"Ticket: {ticket_data['key']}", style='Heading 3').alignment = 1
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", style='Normal').alignment = 1
        
        # Page break
        doc.add_page_break()
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        exec_summary = analysis_results.get('executive_summary', 'No executive summary available.')
        doc.add_paragraph(exec_summary)
        
        # Risk Assessment Summary
        doc.add_heading('Risk Assessment Summary', level=2)
        
        # Create risk summary table
        risk_table = doc.add_table(rows=4, cols=2)
        risk_table.style = 'Table Grid'
        
        risk_data = [
            ('Fraud Risk Score', f"{analysis_results.get('fraud_risk_score', 'Unknown')}/10"),
            ('Security Risk Level', analysis_results.get('security_risk_level', 'Unknown')),
            ('Business Impact', analysis_results.get('business_impact', 'Unknown')),
            ('Analysis Date', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        ]
        
        for i, (field, value) in enumerate(risk_data):
            risk_table.cell(i, 0).text = field
            risk_table.cell(i, 1).text = str(value)
            # Bold the first column
            risk_table.cell(i, 0).paragraphs[0].runs[0].bold = True
        
        # Ticket Information
        doc.add_heading('Ticket Information', level=1)
        
        ticket_table = doc.add_table(rows=10, cols=2)
        ticket_table.style = 'Table Grid'
        
        ticket_info = [
            ('Ticket Key', ticket_data['key']),
            ('Summary', ticket_data['summary'][:100] + '...' if len(ticket_data['summary']) > 100 else ticket_data['summary']),
            ('Type', ticket_data['issue_type']),
            ('Priority', ticket_data['priority']),
            ('Status', ticket_data['status']),
            ('Assignee', ticket_data['assignee']),
            ('Reporter', ticket_data['reporter']),
            ('Created', ticket_data['created'][:10] if ticket_data['created'] else 'N/A'),
            ('Last Updated', ticket_data['updated'][:10] if ticket_data['updated'] else 'N/A'),
            ('Security Level', ticket_data['security_level'])
        ]
        
        for i, (field, value) in enumerate(ticket_info):
            ticket_table.cell(i, 0).text = field
            ticket_table.cell(i, 1).text = str(value)
            ticket_table.cell(i, 0).paragraphs[0].runs[0].bold = True
        
        # Key Findings
        doc.add_heading('Key Findings', level=1)
        findings = analysis_results.get('key_findings', [])
        if findings:
            for finding in findings:
                doc.add_paragraph(finding, style='List Bullet')
        else:
            doc.add_paragraph('No specific findings identified.')
        
        # Red Flags
        if analysis_results.get('red_flags'):
            doc.add_heading('Security Red Flags', level=1)
            for flag in analysis_results['red_flags']:
                p = doc.add_paragraph(flag, style='List Bullet')
                # Make red flags bold and red (if possible)
                for run in p.runs:
                    run.bold = True
        
        # Recommendations
        doc.add_heading('Recommendations', level=1)
        recommendations = analysis_results.get('recommendations', [])
        if recommendations:
            for i, rec in enumerate(recommendations, 1):
                doc.add_paragraph(f"{i}. {rec}", style='List Number')
        else:
            doc.add_paragraph('No specific recommendations at this time.')
        
        # Compliance Issues
        if analysis_results.get('compliance_issues'):
            doc.add_heading('Compliance Concerns', level=1)
            for issue in analysis_results['compliance_issues']:
                doc.add_paragraph(issue, style='List Bullet')
        
        # Detailed Analysis
        if analysis_results.get('detailed_analysis'):
            doc.add_heading('Detailed Analysis', level=1)
            doc.add_paragraph(analysis_results['detailed_analysis'])
        
        # Footer
        doc.add_paragraph('\n' + '='*50)
        footer_info = [
            f"Report generated by: Jira AI Security Analysis System",
            f"Analysis performed on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            f"Ticket analyzed: {ticket_data['key']}",
            f"Report version: 1.0"
        ]
        
        for info in footer_info:
            doc.add_paragraph(info, style='Normal')
        
        # Save to bytes
        doc_stream = BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)
        
        return doc_stream.getvalue()
    
    def process_security_analysis_command(self, user_command: str) -> Dict[str, Any]:
        """Main method to process security analysis commands end-to-end"""
        logger.info(f"Processing security analysis command: '{user_command}'")
        
        try:
            # Parse the command
            parsed_command = self.parse_security_command(user_command)
            
            if not parsed_command['success']:
                return parsed_command
            
            ticket_key = parsed_command['ticket_key']
            
            # Get comprehensive ticket data
            ticket_data = self.get_comprehensive_ticket_data(ticket_key)
            
            if not ticket_data:
                return {
                    'success': False,
                    'error': f'Ticket {ticket_key} not found or inaccessible',
                    'ticket_key': ticket_key
                }
            
            # Perform AI analysis
            analysis_results = self.perform_ai_security_analysis(ticket_data)
            
            # Generate Word document
            doc_bytes = self.generate_security_report_document(ticket_data, analysis_results)
            
            # Create temporary file for download
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            temp_file.write(doc_bytes)
            temp_file.close()
            
            # Generate unique download ID
            download_id = str(uuid.uuid4())
            filename = f'Security_Analysis_{ticket_key}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
            
            return {
                'success': True,
                'ticket_key': ticket_key,
                'analysis_summary': analysis_results.get('executive_summary', 'Analysis completed successfully'),
                'fraud_risk_score': analysis_results.get('fraud_risk_score', 'Unknown'),
                'security_risk_level': analysis_results.get('security_risk_level', 'Unknown'),
                'business_impact': analysis_results.get('business_impact', 'Unknown'),
                'key_findings_count': len(analysis_results.get('key_findings', [])),
                'recommendations_count': len(analysis_results.get('recommendations', [])),
                'document_path': temp_file.name,
                'download_id': download_id,
                'filename': filename,
                'analysis_timestamp': datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Error processing security analysis command: {e}")
            return {
                'success': False,
                'error': str(e),
                'ticket_key': parsed_command.get('ticket_key', 'Unknown')
            }
